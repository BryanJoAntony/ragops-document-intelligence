from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from app.utils.text import normalize_whitespace


@dataclass(frozen=True)
class ParsedPage:
    page_number: int | None
    text: str


@dataclass(frozen=True)
class ParsedDocument:
    parser_name: str
    full_text: str
    pages: list[ParsedPage]


class DocumentParserService:
    def parse(self, filename: str, content_type: str, file_bytes: bytes) -> ParsedDocument:
        suffix = Path(filename).suffix.lower()

        if suffix == ".txt" or content_type.startswith("text/"):
            return self._parse_txt(file_bytes)

        if suffix == ".pdf" or content_type == "application/pdf":
            return self._parse_pdf(file_bytes)

        if suffix == ".docx" or content_type.endswith("wordprocessingml.document"):
            return self._parse_docx(file_bytes)

        raise ValueError(f"Unsupported file type for parsing: {filename}")

    def _parse_txt(self, file_bytes: bytes) -> ParsedDocument:
        text = file_bytes.decode("utf-8", errors="replace")
        text = normalize_whitespace(text)

        return ParsedDocument(
            parser_name="txt_parser",
            full_text=text,
            pages=[ParsedPage(page_number=1, text=text)],
        )

    def _parse_pdf(self, file_bytes: bytes) -> ParsedDocument:
        # Lazy import keeps PyMuPDF/SWIG warnings out of tests that do not parse PDFs.
        import fitz

        pages: list[ParsedPage] = []

        with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
            for index, page in enumerate(pdf, start=1):
                page_text = normalize_whitespace(page.get_text("text"))
                if page_text:
                    pages.append(ParsedPage(page_number=index, text=page_text))

        full_text = normalize_whitespace("\n\n".join(page.text for page in pages))

        return ParsedDocument(
            parser_name="pymupdf_text_parser",
            full_text=full_text,
            pages=pages,
        )

    def _parse_docx(self, file_bytes: bytes) -> ParsedDocument:
        # Lazy import keeps optional parser dependencies isolated to the file types that need them.
        from docx import Document as DocxDocument

        doc = DocxDocument(BytesIO(file_bytes))
        paragraphs = [normalize_whitespace(paragraph.text) for paragraph in doc.paragraphs]
        text = normalize_whitespace("\n\n".join(p for p in paragraphs if p))

        return ParsedDocument(
            parser_name="python_docx_parser",
            full_text=text,
            pages=[ParsedPage(page_number=None, text=text)],
        )